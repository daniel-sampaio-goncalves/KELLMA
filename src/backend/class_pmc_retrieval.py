#Tested with Python 3.14.1

from typing import List, Any
import concurrent.futures
import copy
import json
import multiprocessing
from datetime import datetime
from pathlib import Path


from ollama import GenerateResponse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


import rag_class
import pmc_article_parser_functions


class PmcRetrievalPipeline:
    """This class prases PMC articles and feeds them into a RAG pipeline to refine search results.
    Attributes:
        articles_path (list[str]|list[Path]): path of all the articles to search in
        master_keyword (str): a master keyword that will need to be in the text chunks (any match, case insensitive) eg.'sox2 '
        textsearch_keywords (list[str]):keywords, where at least one of which is be present with the masterkeyword 
                                        in the text chunks eg. ["neuron", "Brain", "Cortex"]
        master_keyword_rag (str): A master keyword used to stengthen the semantics relation in RAG 'Sox2 gene'.
        keywords_for_rag (list[str]): A list of keywords that capture the semantics of what you want to refine using RAG.
                                      eg. ["neuron","cortex","brain","marker for neurons", "neural precursors","progenitor","neuronal lineage marker","expression"]
        threshold_embedding (float, default=0.5): a float 0.0-1 to filter embedding scores before feeding to LLMs
        llm_approvals (int, default=2): Out of 4 different models, how many need to approve the relation between keywords and text in order to retain the final citation
        llm_keywords_for_approval (list[str]): List of two str elements used to feed into the LLMs system prompt to refine the semantics between the first element of the list and text chunks which passed through RAG pipeline. 
            The first element is a word semantic or sentencence you're looking to identify, the sencod should be a description or example to give better context to the llms. 
            Eg: ["Neurons - Markers for neurons - Brain structures",
                "cell type expression, genes related or involved to neuronal processes, neuronal processes (even if broader), brain structures, or any biological process closely related"]
        run_embedding_in_multithreaded_batches (bool, default=False): whether to run embedding in multithreaded batches of articles or not. Default is False.
             If False, the embedding function will run all the the list 'keywords_for_rag' per together in multithreading per paper sequentially instead.
             If True, the embedding function will run all the 'keywords' sequentially but many papers at the same time.  
             **Note: Generally if using a single ollama instance (1 GPU) processing times processing papers in parallel is sligthly faster. But depending how the amount of papers, keywords used and chunk size one method can be faster than the other. 
                     Eg. For a single instance using smaller chunck sizes and few keywords  processing papers sequentially is faster, if many keywords processing papers in paralell is slightly faster.
        number_of_batches_to_run_embedding (int, default=10): Number of max keys in the subdictonnaries will have before being processed (in parallel or sequentially during embedding) after spliting main dictionnary
        suffix_for_saved_docs (str, default=""): Suffix to add to the saved documents. This is useful for organizing your saved documents.
    """
    def __init__(
            self,
            articles_path: list[str]|list[Path]|None=None,
            master_keyword: str|None = None,
            textsearch_keywords: list[str]|None=None,
            master_keyword_rag: str|None = None,
            keywords_for_rag: list[str]|None=None,
            threshold_embedding: float=0.5,
            llm_keywords_for_approval: list[str]= ["Neurons - Markers for neurons - Brain structures",
                                "cell type expression, genes related or involved to neuronal processes, neuronal processes (even if broader), brain structures, or any biological process closely related"],
            llm_approvals: int=2,
            run_embedding_in_multithreaded_batches:str="False",
            number_of_batches_to_run_embedding: int=10,
            suffix_for_saved_docs=""
        ) -> None:

#       #Raise value errors if necessary attributes aren't provided:
        if articles_path is None: 
            raise ValueError("Please provide a list of articles paths")
        else:
            self.articles_path: list[str]|list[Path] =articles_path

#        
        if master_keyword is None:
            raise ValueError('''
                             Please provide a given master keyword that needs to be present in the text
                             eg. sox2
                             ''')
        else:
            self.master_keyword: str = master_keyword        

#
        if textsearch_keywords is None:
            raise ValueError('''
                             Please provide a list keywords to prescreen articles"
                             eg. ["neuron", "Brain", "Cortex"]''')
        else:
            self.textsearch_keywords: list[str] = textsearch_keywords

#
        if master_keyword_rag is None:
            raise ValueError('''
                             Please provide a masterkey word to use in the RAG pipeline eg. "Sox2 gene".
                             ''')
        else:
            self.master_keyword_rag: str=master_keyword_rag

#       
        if keywords_for_rag is None:
            raise ValueError('''
                             Please provide key words for RAG eg.
                             ["neuron","cortex","brain","marker for neurons", "neural precursors","progenitor","neuronal lineage marker","expression"]
                             ''')
        else:
            self.keywords_for_rag:  list[str]=keywords_for_rag
        
        self.run_embedding_in_multithreaded_batches: bool={"True": True, "False": False}[run_embedding_in_multithreaded_batches] #change arguments from text to bool
        self.number_of_batches_to_run_embedding: int= number_of_batches_to_run_embedding
        self.suffix_for_saved_docs: str= suffix_for_saved_docs

        #Threshold metrics:
        self.threshold_embedding: float=threshold_embedding
        self.llm_approvals:       int= llm_approvals 
        self.llm_keywords_for_approval: list[str]=llm_keywords_for_approval
        #attributes to store method results:
        self.parsed_article_list:       dict[str, list[str]]={}
        self.reference_list:            list[str]=[]
        self.reference_chunk_count:     dict[str, int]={}
        self.max_paper_score:           dict[str, float]={} 
        self.max_paper_score_filter:    dict[str, float]={} 
        self.best_chunks_in_per_paper:  dict[str, list[str|float]]={} 
        self.high_interest_chunks:      dict[str, list[tuple[str,float]]]={} 
        self.result_dict:               dict[str, list[tuple[str,float]]]={}  
        
        self.llm_reviewed_text:         dict[str,dict[str,int]]={}
        self.llm_individual_text_score: dict[str,dict[str,int]]={}
        self.llm_approved_dict_:        dict[str, list[str]]={}
        self.llm_AI_response:           dict[str, dict[str,str]]={}
        
        self.LLM_Summary:               Any = {} 
        self.summary_json:             dict[str, list[str]|str]={}


    def process_multiple_articles(self) -> None:
            """Here the function processes a file list. It processes all the articles in parallel and merges all the dictionnaries into a single dictionary with the article id as key and the list of chunks text as value"""
            Final_list_for_rag_: dict[str,list[str]] = {}
            
            max_threads: int=multiprocessing.cpu_count()

            # for compatibility use ThreadPoolExecutor for python <3.14 (eg. 150k articles are parsed in 1min 26s using an i9 10900K)
            #with concurrent.futures.ThreadPoolExecutor(max_workers=max_threads) as executor:
            
            #for python >3.14 keep InterpreterPoolExecutor -> uses true multithreading (eg. 150k articles in 17sec using 20 threads using an i9 10900K)

            with concurrent.futures.InterpreterPoolExecutor(max_workers=max_threads) as executor:  
                futures: list[concurrent.futures.Future[Any]]  = [
                    executor.submit(
                    pmc_article_parser_functions.return_single_pmc_article_processed,
                    self.textsearch_keywords,
                    self.master_keyword,
                    single_article_path
                )
                    for single_article_path in self.articles_path
                ]

                total: int = len(futures)
                
                for i, future in enumerate(concurrent.futures.as_completed(futures), 1):

                
                    result_dict: dict[str,list[str]] = future.result()
                    if "" in result_dict:
                        continue

                    else:  
                        Final_list_for_rag_.update(result_dict)

                    print(f"[{i}/{total}] Finished processing {list(result_dict.keys())}")
            
            self.parsed_article_list = Final_list_for_rag_

    
    def clean_attribute_dict(self, attribute_name: dict[str,list[str]]) -> dict[str,list[str]]:
        '''This function removes any empty values from a dictionnary'''
        #First discard any invalid values in the lists inside the dict
        No_empty_vaues_dict: dict[str, list[str]]={     
        keys: ([word for word in value if word not in ("", [], {},[""],[" "])]) for keys, value in attribute_name.items()}
        #return only non empty lists
        return  {key:value for key, value in No_empty_vaues_dict.items() if value not in ("", [], {},[""],[" "])}

    #def create_list_of_references__(self, paper_dictionnary) -> Any:
    #    """This function takes the initial attribute list and splits transforms the dict into two lists
    #    [0]:Paperref [1]:list with text chunk""" 
    #    self.reference_list= [key for key in paper_dictionnary]
    #    self.reference_chunk_count={key:len(paper_dictionnary[key]) for key in paper_dictionnary}


    def create_list_of_references(self, paper_dictionnary:dict[str, List[str]]) -> Any:
        """This function takes the initial attribute list and splits transforms the dict into two lists
        [0]:Paperref [1]:list with text chunk""" 
        reference_list: list[str]= [key for key in paper_dictionnary]
        reference_chunk_count: dict[str,int]={key:len(paper_dictionnary[key]) for key in paper_dictionnary}
        return reference_list, reference_chunk_count
    
    def prepare_list_to_run(self, paper_dict:dict[str, List[str]], print_subprocess_embeddings:bool=False) -> Any:
        """This function takes a dictionary of papers, runs embedding to compare semantics and returns multiple dictionnaries of filtered and unfiltered embedding results."""

        if self.run_embedding_in_multithreaded_batches:
            multithreading=False
        else:
            multithreading=True
        paper_dic_cleaned: dict[str, List[str]] =self.clean_attribute_dict(paper_dict)
        
        local_reference_list: list[str]
        reference_chunk_count: dict[str,int]
        local_reference_list,reference_chunk_count=self.create_list_of_references(paper_dict)
        
        local_result_dict:dict[str, list[tuple[str,float]]]={}

        
        #shallow copy and add master keyword
        rag_keywords_to_compare_to: list[str] = [*self.keywords_for_rag, self.master_keyword_rag]
        

        for index_ in range(len(local_reference_list)):

            if print_subprocess_embeddings:
                print(f'Processing {local_reference_list[index_]} {index_+1}/{len(local_reference_list)}')
            rag_pipeline=rag_class.Refined_Semantic_matching(
            array_word_to_check_for_embeddings=rag_keywords_to_compare_to,
            vector_items_to_compare=paper_dic_cleaned[local_reference_list[index_]],
            adjust_for_rank = True,
            embedding_rank_percent_cutoff = self.threshold_embedding,
            embedding_cpu_threadding=multithreading
            )
            rag_pipeline.run_embedding_scores()
            
            local_result_dict.update({local_reference_list[index_]:rag_pipeline.averaged_similarity_results})
        
        

        paper_scores: list[list[float]]=[[item[1] for item in local_result_dict[key]] for key in local_result_dict]
        
        local_max_paper_score: dict[str, float] ={local_reference_list[item[0]]:max(item[1]) for item in enumerate(paper_scores)}
        
        
        #filter if max score > threshold_embedding

        local_max_paper_score_filter: dict[str, float] ={key:value for key, value in local_max_paper_score.items() if value > self.threshold_embedding}
        
        #Chunk:
       
        local_best_chunks_in_per_paper:dict[str, list[str|float]] ={key:[item[0],item[1]] for key in local_result_dict for item in local_result_dict[key] if item[1] == local_max_paper_score_filter.get(key)}
        
        #Chunk:
        
        local_high_interest_chunks: dict[str, list[tuple[str, float]]]={key: [item for item in local_result_dict.get(key,  []) if item[1] > self.threshold_embedding] for key in local_result_dict }
        #remove empty keys / keeps only truthy lists (non-empty)
        local_high_interest_chunks = { key: value for key, value in local_high_interest_chunks.items()  if value }

        return local_max_paper_score, local_max_paper_score_filter, local_best_chunks_in_per_paper, local_high_interest_chunks, local_result_dict, local_reference_list, reference_chunk_count
         
    def run_in_batches(self) -> None:
        '''This function takes splits the main dictionnary into a list of subdictionnaries to runs in batches for speed and memory efficiency.
           Generally processing papers in parallel is faster compared to sequentially, but depending on GPU performance and if using many keywords processing keywords in parallel can be faster.
           Generally if using a single ollama instance (1 GPU) processing times processing keywords in parallel is sligthly faster. But depending how the amount of papers, keywords used and chunk size one process can be faster than the other. 
           Note**: The function could be optimized for multithreading if using multiple ollama instances listening on multiple ports with multiple GPUs eg sending a batch to a different port.
        '''
        #If no cached list run the parsing func again
        if not self.parsed_article_list:
            self.process_multiple_articles()
            
            #if the list returns empty exit 
            if not self.parsed_article_list: 
                print("No articles found. Exiting.", flush=True) 
                raise RuntimeError("No articles found matching the keyword search")
                #sys.exit(1)


        chunks_of_dicts_values: list[tuple[str, list[str]]] = list(self.parsed_article_list.items())
        chunk_size: int = self.number_of_batches_to_run_embedding
        batch_dicts: list[dict[str, list[str]]] = [dict(chunks_of_dicts_values[i:i+chunk_size]) for i in range(0, len(chunks_of_dicts_values), chunk_size)]
  
        embedding_results: list[Any] = []
        completed = 0
        
        if self.run_embedding_in_multithreaded_batches:
            max_cpu_workers: int= multiprocessing.cpu_count()
            print(f"Using run in batches in multithreaded mode -- workers: {max_cpu_workers}")
            with concurrent.futures.ThreadPoolExecutor(max_workers=max_cpu_workers) as executor:
                futures: List[concurrent.futures.Future[Any]] = [executor.submit(self.prepare_list_to_run, batch) for batch in batch_dicts]
                total_tasks: int = len(futures)
                print(f"##### Started embeddings: {total_tasks} total tasks #####")
                # Collect results as they complete
                for future in concurrent.futures.as_completed(futures):
                    result = future.result()
                    embedding_results.append(result)
                    
                    completed += 1
                    remaining: int = total_tasks - completed
                    print(f"##### {completed}/{total_tasks} --> {remaining} tasks left. #####")
                    #except Exception as e:
                        # Handle exceptions gracefully
                        #print(f"Error processing batch: {e}")
        else:
            print(f"""Using run in batches in sequentiall mode
                  ##### Started embeddings: {len(batch_dicts)} total tasks with each {chunk_size} papers #####""")
            embedding_results=[self.prepare_list_to_run(batch,print_subprocess_embeddings=True) for batch in batch_dicts]
            # Note in for debugging, test full verison: remouved this : print_subprocess_embeddings=True) for batch in batch_dicts[0:20]]
        
        
        local_ref_list: list[str]=[]
        for max_paper_score, max_paper_score_filter,best_chunks_in_per_paper, high_interest_chunks, result_dict, reference_list, reference_chunk_count_ in embedding_results:
            self.max_paper_score.update(max_paper_score)
            self.max_paper_score_filter.update(max_paper_score_filter)
            self.best_chunks_in_per_paper.update(best_chunks_in_per_paper)
            self.high_interest_chunks.update(high_interest_chunks)
            self.result_dict.update(result_dict)
            local_ref_list.append(reference_list)
            self.reference_chunk_count.update(reference_chunk_count_)
        
        self.reference_list=[item for key in local_ref_list for item in key]
        print("##### Finished RAG #####")
    
    def run_LLM(self, new_threshold: float|None=None) -> None:
        """This function calls if needed the RAG pipeline, and uses its output to perform a double check using several (currently 4) llm chatbots
           The gpt models will be launched and process every paper sequentially to avoid loading models multiple times
           The models output will be evaluated and text chunks will be stored into a new dictionnary if combined model score > llm_approvals   
        """

        if not self.max_paper_score:
            self.run_in_batches()
        
        elif new_threshold:

            local_high_interest_chunks: dict[str, list[tuple[str,float]]] ={key: [item for item in self.result_dict.get(key,  []) if item[1] > new_threshold] for key in self.result_dict }
            #remove empty keys / keeps only truthy lists (non-empty)
            local_high_interest_chunks: dict[str, list[tuple[str,float]]] = { key: value for key, value in local_high_interest_chunks.items()  if value }
            self.high_interest_chunks=local_high_interest_chunks

        llm_chatting=rag_class.Refined_Semantic_matching(
                array_word_to_check_for_embeddings=self.keywords_for_rag,
                checking_prompt_for_LLMs = self.llm_keywords_for_approval)
        #llm_lunc.model_list
        
        results_semantics: dict[str,dict[str,int]]={}
        model_results: dict[str,dict[str,int]]={}

        final_results_semantics: Any={item:{} for item in self.high_interest_chunks}
        print(final_results_semantics)
        final_model_results: Any={}

        # gather all the LLMs response to a single dictionnary
        for model in llm_chatting.model_list:

            for paper in self.high_interest_chunks:
                
                for chunk in self.high_interest_chunks[paper]:
                    chunk_str:str=chunk[0]
                    #print(f"{paper} and text: {chunk_str}\n")
                    results_semantics, model_results =llm_chatting.Check_LLM_Semantics_sequentially(text_chunks=chunk_str, reference=paper, model_to_use=model)
                    AI_response: str=llm_chatting.LLM_last_AI_response
                    
                    #update the results dictionary with the new results
                    if results_semantics[paper][chunk_str]==1 and chunk_str in final_results_semantics[paper]:
                        final_results_semantics[paper].update({chunk_str:final_results_semantics[paper][chunk_str]+1})
                    else:
                        final_results_semantics[paper].update({chunk_str:results_semantics[paper][chunk_str]})
                    
                    #The same but with the model references
                    if chunk_str not in final_model_results:
                            final_model_results.update({chunk_str:model_results[chunk_str]})
                    else:
                            final_model_results[chunk_str].update({model:model_results[chunk_str][model]})
                    
                    if chunk_str not in self.llm_AI_response:
                        self.llm_AI_response.update({chunk_str:{model:AI_response}})
                    else:
                        self.llm_AI_response[chunk_str].update({model:AI_response})
        
        #remove low scoring text chunks and create a list of text chunks per reference
        self.llm_approved_dict_ ={paper:[chunk for chunk in final_results_semantics[paper] if final_results_semantics[paper][chunk]>self.llm_approvals] for paper in final_results_semantics}
        #remove empty keys:
        self.llm_approved_dict_ ={key:value for key, value in self.llm_approved_dict_.items() if value not in ("", [], [""],[" "])}
        self.llm_reviewed_text = final_results_semantics
        self.llm_individual_text_score = final_model_results
    
    
    def summarize_papers(self, output_word:bool=True, json_file:bool=True, new_threshold:bool=False) -> None:
        """This function runs all the previous methods to prase and retrieve information from articles and summarizes them by calling a LLM.
           It then saves a json style dictionnary with the summary and the llm approved text chunks used to generate the summary
           Depending on the options it also dumps a json or formated word.docx file
        """
        
        if new_threshold:
            self.run_LLM(new_threshold=self.threshold_embedding)

        elif not self.llm_approved_dict_:
            self.run_LLM()

        #Log files:
        result_dir = Path("../../results")
        result_dir.mkdir(exist_ok=True)
        
        #unzip text out of the dicts&lists
        text_prompt_list: list[str]=[text for paper in self.llm_approved_dict_.values() for text in paper]
        total_approved_text_chucks: int=len(text_prompt_list)
        total_approved_papers: int=len(self.llm_approved_dict_.keys())
        #limit the amount of text to 100 text cuncks that is fed into the summarry LLM 
        #this avoids out of memory bugs and poor llm accuracy
        if  len(text_prompt_list) > 100:
            #filter first all the scores out
            retained_thresholds: List[float]=[item2[1] 
            for item in self.llm_approved_dict_.keys()
            for item2 in self.high_interest_chunks[item]]
            #sort and  cut after 100 (not more than 180 entries (<65k tokens) to avoid out of memory) 
            retained_thresholds.sort(reverse=True)
            
            #[0:100] limit the context feed into the LLM to 100 text chunks
            new_scores: List[float]=retained_thresholds[0:100]             
            #recreate the dictionnary containing only the text chunks with the highest scores
            cut_down_list: dict[str, list[str]]={
                key: [item[0] for item in items if item[1] in new_scores]
                for key, items in self.high_interest_chunks.items()
                if any(item[1] in new_scores for item in items)
            }
            #recreate the text prompt list from earlier with the new cut down cuncks
            text_prompt_list: list[str]=[text for paper in cut_down_list.values() for text in paper]
            print("""
                  ###LLMs approved too many text cuncks for summary, llm approved text cuncks will be 
                  cut down to the 100 most highest scoring cuncks during RAG
                  """)

        # join list into a str to feed into llm
        text_prompt: str ="".join(f" Text{key+1}: {text} |" for key,text in enumerate(text_prompt_list))
 
        llm_chatting=rag_class.Refined_Semantic_matching(
            array_word_to_check_for_embeddings=self.keywords_for_rag,
            checking_prompt_for_LLMs = self.llm_keywords_for_approval)
        
        summary_response: GenerateResponse=llm_chatting.summarizing_retrieved_references(text_prompt, gene_search=self.master_keyword_rag)

        self.LLM_Summary=summary_response
        self.summary_json={"Gene name":self.master_keyword}
        self.summary_json.update({"Date": f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'})
        self.summary_json.update({"RAG master keyword": self.master_keyword_rag})
        self.summary_json.update({"Textseach keywords": self.textsearch_keywords})
        self.summary_json.update({"RAG master keyword:":self.master_keyword})
        self.summary_json.update({"Search Criteria RAG":self.master_keyword_rag})
        self.summary_json.update({"Search Criteria LLM":self.llm_keywords_for_approval[0]})
        self.summary_json.update({"LLM context":self.llm_keywords_for_approval[1]})
        self.summary_json.update({"Summary":summary_response.response})
        self.summary_json.update({"RAG threshold used: ":str(self.threshold_embedding)})

        ref_dict: dict[str, list[str]]=copy.deepcopy(self.llm_approved_dict_)
        
        if ref_dict:
            self.summary_json.update(ref_dict)
        else:
            self.summary_json.update({"REFERENCES":"None found"})
        
        if json_file:
            with open(result_dir/f"llm_rag_{self.master_keyword}_llma_{self.llm_approvals}_{self.suffix_for_saved_docs}.json", "w") as f:
                json.dump(self.summary_json, f, indent=1)

        if output_word:
            doc: Any = Document()
            word_title: Any=doc.add_paragraph()
            run_word_title: Any=word_title.add_run(f"RAG Results")
            run_word_title.bold = True
            run_word_title.font.size = Pt(13)
            
            #Add blank
            doc.add_paragraph()
            #add time:
            stamp_time: Any=doc.add_paragraph()
            run_stamp_time: Any=stamp_time.add_run(f"Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}")
            run_stamp_time.bold = False
            run_stamp_time.font.size = Pt(11)
            
            
            #Add table
            table = doc.add_table(rows=11, cols=2)
            table.cell(0, 0).text = "Gene name:"
            table.cell(0, 1).text = f"{self.master_keyword}"
            table.cell(1, 0).text = "Textseach keywords:"
            table.cell(1, 1).text = f"{self.textsearch_keywords}"
            table.cell(2, 0).text = "RAG master keyword:"
            table.cell(2, 1).text = f"{self.master_keyword_rag}"
            table.cell(3, 0).text = "Search Criteria RAG:"
            table.cell(3, 1).text = f"{self.keywords_for_rag}"
            table.cell(4, 0).text = "Search Criteria LLM"
            table.cell(4, 1).text = f"{self.llm_keywords_for_approval[0]}"
            table.cell(5, 0).text = "LLM context:"
            table.cell(5, 1).text = f"{self.llm_keywords_for_approval[1]}"
            table.cell(6, 0).text = f"RAG threshold used: "
            table.cell(6, 1).text = f"{self.threshold_embedding}"
            table.cell(7, 0).text = f"Total papers scanned: "
            table.cell(7, 1).text = f"{len(self.articles_path)}"
            table.cell(8, 0).text = f"Total papers found: "
            table.cell(8, 1).text = f"{len(self.parsed_article_list)}"
            table.cell(9, 0).text = f"Total approved papers: "
            table.cell(9, 1).text = f"{total_approved_papers}"
            table.cell(10, 0).text = f"Total approved text cuncks: "
            table.cell(10, 1).text = f"{total_approved_text_chucks}"

            #Add blank
            doc.add_paragraph()

            paragraph_title=doc.add_paragraph()
            run_paragraph_title=paragraph_title.add_run(f"Summary: ")
            run_paragraph_title.bold = True
            run_paragraph_title.font.size = Pt(11)

            
            paragraph_summary=doc.add_paragraph()
            paragraph_summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_paragraph_summary=paragraph_summary.add_run(f"{summary_response.response}")
            run_paragraph_summary.bold = False
            run_paragraph_summary.font.size = Pt(11)

            #add references to the word document, inside a table; concat several strings per ref
            doc.add_paragraph()

            if ref_dict:
                n_max_rows: int=len(ref_dict)
                table_refs: Any = doc.add_table(rows=n_max_rows+1, cols=2)
                

                for idx, key in enumerate(ref_dict):
                    
                    table_refs.cell(idx, 0).text = key
                    #add a seperator first between each ref text
                    values: list[str] = ref_dict[key]
                    formatted = "[...]".join(values)
                    table_refs.cell(idx, 1).text = formatted


            else:
                table_refs: Any = doc.add_table(rows=1, cols=2)
                table_refs.cell(0, 0).text= "NO REFERENCE FOUND"
                table_refs.cell(0, 1).text= ""
            
            for row in table_refs.rows: 
                    cell = row.cells[1] 
                    for parag in cell.paragraphs: 
                        parag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            #disable autofit, and fit refs and results
            table_refs.autofit = False 
            table_refs.allow_autofit = False

            # enforce widths per cell 
            for row in table_refs.rows: 
                row.cells[0].width = Cm(2.9) 
                row.cells[1].width = Cm(12.7)

            doc.save(result_dir/f"llm_rag_{self.master_keyword}_llma_{self.llm_approvals}_{self.suffix_for_saved_docs}.docx")
